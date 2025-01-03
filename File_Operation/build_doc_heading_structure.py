import pandas as pd
from docx import Document


def main():
    # Read the Excel file
    df = pd.read_excel(excel_file)

    # Open a new Word document
    doc = Document()

    # Initialize the last heading of each level
    last_heading = {1: None}

    # Iterate over the DataFrame
    for index, row in df.iterrows():
        for column in df.columns:
            # Ignore NaN values
            if pd.notna(row[column]):
                # Get the current heading level
                current_heading_level = int(column)

                # If the current heading level is higher than the last heading level,
                # add a new heading under the last heading of the previous level
                if current_heading_level > max(last_heading.keys()):
                    last_heading[current_heading_level] = doc.add_heading(row[column], level=current_heading_level)
                # If the current heading level is the same as the last heading level,
                # add a new heading under the last heading of the same level
                elif current_heading_level == max(last_heading.keys()):
                    last_heading[current_heading_level] = doc.add_heading(row[column], level=current_heading_level)
                # If the current heading level is lower than the last heading level,
                # add a new heading under the last heading of its level
                else:
                    last_heading[current_heading_level] = doc.add_heading(row[column], level=current_heading_level)
                    # Remove any higher level headings from the last_heading dictionary
                    for key in list(last_heading.keys()):
                        if key > current_heading_level:
                            del last_heading[key]

                # Add blank lines or page break after the heading
                # doc.add_paragraph()  # First blank line
                # doc.add_paragraph()  # Second blank line
                doc.add_page_break()  # Page Break

    # Save the Word document
    doc.save(doc_file)


if __name__ == "__main__":
    excel_file = (r"D:\YK\Honeywell\2024 MYP-000360 Biocon BMS & EMS HMI Replacement\1 Project Document\3.0 Project Engineering\GRAPHICS APPROVAL "
                  r"PACKAGE\BMS\v0.6\HMI BMS TOC.xlsx")
    doc_file = r"D:\YK\Downloads\1.docx"
    main()
